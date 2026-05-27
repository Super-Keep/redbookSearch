#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
@Time    : 2026-05-23
@Author  : AI Assistant
@File    : word_export_util.py
@Desc    : Export AI analysis to Word (.docx) document with markdown-like formatting
"""
import os
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.klogger_util import logger


# Default export directory
_DEFAULT_EXPORT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "exports"
)


def export_analysis_to_word(
    title: str,
    content: str,
    keyword: str = "",
    note_count: int = 0,
    export_dir: str = ""
) -> str:
    """
    Export AI analysis content (markdown) to a Word document.

    :param title: Document title
    :param content: Markdown formatted AI analysis content
    :param keyword: Search keyword (for metadata)
    :param note_count: Number of notes analyzed
    :param export_dir: Directory to save the file (defaults to ./exports/)
    :return: Full path to the saved .docx file
    """
    export_dir = export_dir or _DEFAULT_EXPORT_DIR
    os.makedirs(export_dir, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_text = f"关键词: {keyword} | 分析笔记数: {note_count} | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    meta_para = doc.add_paragraph(meta_text)
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")  # Spacer

    # Parse markdown content and add to document
    _markdown_to_docx(doc, content)

    # Generate filename
    safe_keyword = re.sub(r'[\\/:*?"<>|\n\r]+', '', keyword)[:20].strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"小红书分析_{safe_keyword}_{timestamp}.docx"
    filepath = os.path.join(export_dir, filename)

    doc.save(filepath)
    logger.info(f"Word document exported: {filepath}")

    return filepath


def _markdown_to_docx(doc: Document, markdown: str) -> None:
    """
    Convert markdown text to Word document content.
    Supports: headings, bold, bullet lists, ordered lists, paragraphs.

    :param doc: python-docx Document instance
    :param markdown: Markdown content string
    """
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Heading 1: # Title
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=1)

        # Heading 2: ## Title
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)

        # Heading 3: ### Title
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)

        # Heading 4: #### Title
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)

        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 50)

        # Bullet list: - item or * item
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            _add_rich_text(para, text)

        # Numbered list: 1. item
        elif len(line.strip()) > 2 and line.strip()[0].isdigit() and ". " in line.strip()[:5]:
            text = line.strip().split(". ", 1)[1] if ". " in line.strip() else line.strip()
            para = doc.add_paragraph(style='List Number')
            _add_rich_text(para, text)

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_rich_text(para, line.strip())

        i += 1


def _add_rich_text(paragraph, text: str) -> None:
    """
    Add text with inline formatting (bold, italic) to a paragraph.

    :param paragraph: python-docx Paragraph instance
    :param text: Text with possible **bold** or *italic* markers
    """
    # Split by bold markers **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
